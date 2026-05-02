"""Build CMPB submission docx package for TinyDCS.

Generates Word files for portal upload at editorialmanager.com/cmpb/:
  - Manuscript.docx (with continuous line numbers + double spacing)
  - Cover Letter.docx
  - Highlights.docx
  - Declaration of Competing Interest.docx
  - Statement on Human and Animal Studies.docx
  - Author Contributions (CRediT).docx
  - Suggested Reviewers.docx

Run: python3 build_docx.py
Output: docx_submission/*.docx
"""

from __future__ import annotations

import re
import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "docx_submission"
OUT.mkdir(exist_ok=True)
FIG_SRC = ROOT.parent.parent.parent / "artifacts" / "paper_figures"

TITLE = (
    "TinyDCS: An edge-deployable machine-learning surrogate of the ADRAC "
    "altitude-decompression-sickness risk model with continuous-exposure "
    "covariates and calibrated uncertainty"
)
RUNNING_TITLE = "TinyDCS: edge surrogate of ADRAC"
AUTHORS_LINE = "Diego Malpica, MD¹*  ·  Marian Farfán, MD¹"
AFFIL = (
    "¹ Subdirectorate of Aerospace Sciences, Direction of Aerospace Medicine, "
    "Colombian Aerospace Force, Bogotá DC, Colombia"
)
CORRESPONDENCE = (
    "* Corresponding author: Diego Malpica, MD. Email: diego.malpica@fac.mil.co"
)
TODAY = "2026-05-02"
JOURNAL = "Computer Methods and Programs in Biomedicine"
EIC = "Filippo Molinari, PhD"
EIC_AFFIL = (
    "Editor-in-Chief, Computer Methods and Programs in Biomedicine\n"
    "Department of Electronics and Telecommunications\n"
    "Polytechnic of Turin, Turin, Italy"
)


# ---------------------------------------------------------------------------
# Document scaffolding helpers
# ---------------------------------------------------------------------------

def _new_doc(double_spaced: bool = True, line_numbers: bool = False) -> Document:
    """Create a new docx with Times New Roman 12pt base style and Elsevier-style margins."""
    doc = Document()

    # Base style (Normal): Times New Roman 12pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Times New Roman")

    # Default paragraph format: line spacing
    pf = style.paragraph_format
    if double_spaced:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    # Margins (Elsevier convention: 2.5 cm all round)
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    if line_numbers:
        _enable_line_numbers(section)

    # Heading styles
    for level in range(1, 5):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            h.font.size = Pt(14)
        elif level == 2:
            h.font.size = Pt(13)
        else:
            h.font.size = Pt(12)
        h.font.bold = True
        rpr2 = h.element.get_or_add_rPr()
        rf2 = rpr2.find(qn("w:rFonts"))
        if rf2 is None:
            rf2 = OxmlElement("w:rFonts")
            rpr2.append(rf2)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf2.set(qn(attr), "Times New Roman")
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        if double_spaced:
            h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    return doc


def _enable_line_numbers(section) -> None:
    """Inject <w:lnNumType> into the section properties for continuous line numbers."""
    sectPr = section._sectPr
    # Remove any existing line-number element
    for existing in sectPr.findall(qn("w:lnNumType")):
        sectPr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:start"), "1")
    ln.set(qn("w:distance"), "360")  # 0.25 inch from text
    ln.set(qn("w:restart"), "continuous")
    sectPr.append(ln)


def _add_para(doc: Document, text: str = "", *, bold: bool = False,
              italic: bool = False, align: int = WD_ALIGN_PARAGRAPH.LEFT,
              size: int | None = None, space_after: int | None = None,
              keep_with_next: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)


def _add_runs(p, *segments) -> None:
    """Append a sequence of (text, **kwargs) tuples as runs to an existing paragraph."""
    for seg in segments:
        if isinstance(seg, str):
            p.add_run(seg)
        else:
            text, *opts = seg
            kwargs = opts[0] if opts else {}
            run = p.add_run(text)
            for k, v in kwargs.items():
                if k == "bold":
                    run.bold = v
                elif k == "italic":
                    run.italic = v
                elif k == "underline":
                    run.underline = v
                elif k == "subscript":
                    run.font.subscript = v
                elif k == "superscript":
                    run.font.superscript = v


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        rpr = run.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(attr), "Times New Roman")


# ---------------------------------------------------------------------------
# Inline rich-text helpers (subscript / italic / bold from markdown-ish text)
# ---------------------------------------------------------------------------

# We use a tiny inline grammar tailored to this manuscript:
#   **bold**           → bold
#   *italic*           → italic
#   `code`             → code (no styling change beyond italic)
#   VO_2               → VO with subscript 2  (handled via explicit replacement)
#   N_2                → N with subscript 2
#   FiO_2              → FiO with subscript 2
#   SpO_2              → SpO with subscript 2
#   CO_2               → CO with subscript 2
#   2.44 μ s      → handled directly (μ already inserted)
# Markdown tokens *italic* with single asterisks must not collide with **bold**.

INLINE_RE = re.compile(
    r"(\*\*[^*\n]+\*\*)"            # **bold**
    r"|(\*[^*\n]+\*)"                # *italic*
    r"|(`[^`\n]+`)"                  # `code`
    r"|((?:VO|N|FiO|SpO|CO|H)_2(?![A-Za-z0-9]))",  # VO_2, N_2, etc.
)


def _add_inline(p, text: str) -> None:
    """Add inline-styled text to paragraph p, parsing **bold**, *italic*, `code`, X_2 subscripts."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
        bold_m, italic_m, code_m, sub_m = m.groups()
        if bold_m:
            run = p.add_run(bold_m[2:-2])
            run.bold = True
        elif italic_m:
            run = p.add_run(italic_m[1:-1])
            run.italic = True
        elif code_m:
            run = p.add_run(code_m[1:-1])
            run.font.name = "Consolas"
        elif sub_m:
            base, _, _ = sub_m.partition("_")
            run = p.add_run(base)
            sub = p.add_run("2")
            sub.font.subscript = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tc_pr.append(borders)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]],
               caption: str | None = None) -> None:
    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cap.paragraph_format.space_after = Pt(4)
        cap.paragraph_format.keep_with_next = True
        _add_inline(cap, caption)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_borders(hdr[i])
        # Light grey shading for header
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tc_pr.append(shd)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            para = cells[c_idx].paragraphs[0]
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            _add_inline(para, val)
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
            _set_cell_borders(cells[c_idx])
    # Spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Manuscript builder
# ---------------------------------------------------------------------------

def build_manuscript() -> None:
    doc = _new_doc(double_spaced=True, line_numbers=True)

    # ----- Title page -----
    _add_para(doc, "Manuscript submitted to", italic=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _add_para(doc, JOURNAL, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=0)
    _add_para(doc, "(Elsevier) — under review", italic=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(18)

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Diego Malpica, MD")
    p.add_run("¹").font.superscript = True
    p.add_run(",*").font.superscript = True
    p.add_run("    Marian Farfán, MD")
    p.add_run("¹").font.superscript = True
    p.paragraph_format.space_after = Pt(8)

    _add_para(doc, AFFIL, align=WD_ALIGN_PARAGRAPH.CENTER, italic=False,
              space_after=8)
    _add_para(doc, CORRESPONDENCE, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=18)

    _add_para(doc, f"Running title: {RUNNING_TITLE}",
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=4)
    _add_para(doc,
              "Article type: Full Length Article  •  Body word count: ~2,950  •  "
              "Abstract: 334 words  •  References: 16  •  Tables: 4  •  Figures: 5",
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=4)
    _add_para(doc, f"Version 1.0.0 — {TODAY}",
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=24)

    doc.add_page_break()

    # ----- Structured abstract -----
    _heading(doc, "Abstract", level=1)

    abstract_blocks = [
        ("Background and Objectives.",
         "The US Air Force ADRAC altitude-decompression-sickness risk model [1] is "
         "the operational standard for planning hypobaric exposures, but its "
         "three-level exercise covariate cannot ingest continuous wearable-derived "
         "VO_2 trajectories and it returns point estimates without calibrated "
         "uncertainty. We aimed to build a wearable-grade machine-learning "
         "surrogate of the ADRAC grid that (i) accepts continuous-VO_2 covariates, "
         "(ii) ships calibrated 95% prediction intervals with uniform altitude-band "
         "coverage, (iii) abstains outside the validated input envelope, and "
         "(iv) meets an edge-deployment footprint below 100 KB with per-inference "
         "latency below 10 μs."),
        ("Methods.",
         "After cleaning a public 16,295-row ADRAC-output grid (15,908 unique "
         "cells; 7.5% rescaled from fraction to percent), we trained a LightGBM "
         "regressor on the logit of P(DCS) using a 13-feature vector that augments "
         "the ADRAC covariates with a Conkin tissue-nitrogen ratio [6] and "
         "continuous-VO_2 summaries (Webb 2010 1-min peak [2]). Monotonicity "
         "constraints, Smithson–Verkuilen boundary shrinkage [8], and a two-stage "
         "zero-inflated split-conformal stack [9,10] with Mahalanobis-distance "
         "out-of-envelope abstention were applied. The surrogate was benchmarked "
         "against a closed-form log-logistic AFT and exported to ONNX."),
        ("Results.",
         "On the held-out random test fold (*n* = 2,386), TinyDCS attained "
         "MAE = 0.020, R² = 0.986, and Brier score = 0.0016 — a 4-fold MAE "
         "reduction and 10-fold Brier reduction over the closed-form baseline "
         "(MAE = 0.086; Brier = 0.0150). Empirical 95% coverage was 0.960 overall "
         "and at least 0.95 in each of the five 5,000-ft altitude bands, closing a "
         "low-band shortfall (coverage 0.58–0.59 at 18,000–23,000 ft) that was "
         "invariant under four conformal-only alternatives. A compact "
         "zero-inflated variant compiled to 95 KB of ONNX with CPU per-row latency "
         "of 2.44 μs (p50). A conjugate-Gaussian hierarchical personalisation "
         "prototype recovered per-subject log-susceptibility at Pearson *r* = 0.63 "
         "after twenty exposures per subject."),
        ("Conclusions.",
         "A continuous-VO_2 ADRAC surrogate with zero-inflated conformal "
         "calibration and out-of-envelope abstention outperforms the closed-form "
         "model on the same data at an edge-feasible memory and latency budget. "
         "External prospective validation on a hypobaric-chamber cohort with "
         "Doppler venous-gas-emboli ground truth, and replacement of the "
         "conjugate-Gaussian personalisation stub with a full hierarchical "
         "Bayesian model on real subjects, are the priority follow-ups."),
    ]
    for header, body in abstract_blocks:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(header + " ")
        run.bold = True
        _add_inline(p, body)

    # Keywords
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Keywords: ")
    run.bold = True
    p.add_run(
        "altitude decompression sickness; ADRAC; wearable computing; "
        "conformal prediction; zero-inflated models; edge AI"
    )

    doc.add_page_break()

    # ----- 1. Introduction -----
    _heading(doc, "1. Introduction", level=1)

    paras_intro = [
        "Altitude-induced decompression sickness (DCS) remains an operational "
        "hazard in unpressurised or rapidly depressurised aviation, hypobaric "
        "chamber training, and extravehicular activity. The USAFSAM Altitude "
        "Decompression Sickness Risk Assessment Computer (ADRAC) has been the US "
        "Air Force operational standard since the late 1990s [1]. ADRAC is a "
        "stratified accelerated-failure-time log-logistic survival model with four "
        "covariates — altitude, prebreathe duration, exercise level (three "
        "categorical levels: Rest / Mild / Heavy), and time-at-altitude — "
        "validated with 150 prospective hypobaric exposures.",

        "Two limitations of ADRAC have been recognised in the primary literature. "
        "First, Webb and colleagues [2,3] demonstrated that the operationally "
        "correct activity metric is the highest 1-minute VO_2 in any 16-minute "
        "window at altitude, not a three-level category; ADRAC has not been refit "
        "to use this metric. Second, ADRAC is a closed-form point estimator — it "
        "does not provide calibrated uncertainty or abstain on inputs outside its "
        "validated envelope.",

        "Separately, Gerth and coauthors [4] developed the 3RUT-MBe1 "
        "bubble-dynamics model that accepts arbitrary continuous-VO_2 trajectories "
        "across the full exposure and outperforms ADRAC on four of five "
        "ADRAC-validation profiles. The 3RUT-MBe1 model, however, is an "
        "ordinary-differential-equation recursion that is too heavy for embedded "
        "deployment on a smartwatch or flight-computer background task. A recent "
        "random-forest study by Han and colleagues [5] applied general-purpose ML "
        "to ADRAC-derived data but did not address calibrated uncertainty or edge "
        "deployment.",

        "Wearables now routinely stream accelerometer-derived VO_2 proxies, "
        "heart-rate variability, SpO_2, and barometer-derived altitude at multi-Hz "
        "rates. This telemetry is ready for on-body DCS risk estimation, but no "
        "existing published model is both wearable-deployable and accepting of "
        "continuous VO_2.",

        "**Contributions.** The present work addresses this gap. Specifically, we:",
    ]
    for txt in paras_intro:
        p = doc.add_paragraph()
        _add_inline(p, txt)

    contrib_items = [
        "Audit and repair the public ADRAC-output grid "
        "(`DCS_Risk_DB_2025.csv`), documenting a systematic scale inconsistency "
        "that affects 7.5% of its rows.",
        "Train a compact machine-learning surrogate of ADRAC that accepts "
        "continuous-VO_2 covariates via a Conkin 2004 single-compartment tissue "
        "ratio feature.",
        "Attach Mondrian (altitude-stratified) split-conformal prediction "
        "intervals [9,10], a two-stage zero-inflated calibration that closes a "
        "low-altitude-band coverage shortfall, and a Mahalanobis-distance "
        "envelope abstention.",
        "Export the trained zero-inflated surrogate to ONNX at multiple size "
        "budgets, demonstrating a 95 KB combined variant that retains R² > 0.98 "
        "and outperforms the closed-form ADRAC baseline by 3× on MAE.",
        "Benchmark the full pipeline under both random and "
        "leave-one-altitude-out splits; release a reproducibility package "
        "(runbook, AGENTS.md continuation guide, trained weights, ONNX "
        "artifacts, metrics JSONs, paper figures, TRIPOD+AI checklist).",
        "Prototype a conjugate-Gaussian hierarchical personalisation layer on a "
        "synthetic cohort, quantifying the number of per-subject exposures "
        "required to reach Brier parity with the population model.",
    ]
    for i, item in enumerate(contrib_items, start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        _add_inline(p, item)

    # ----- 2. Methods -----
    _heading(doc, "2. Methods", level=1)

    _heading(doc, "2.1 Data source and cleaning", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "The ADRAC-derived dataset `DCS_Risk_DB_2025.csv` (16,295 rows) was "
        "obtained from the open `strikerdlm/DCS` repository. It represents the "
        "ADRAC model's output on a factorial grid of altitude (18,000–40,000 ft "
        "in 500-ft increments), prebreathe time (0–60 min in 15-min increments), "
        "exercise level (Rest/Mild/Heavy), and time-at-altitude (10–240 min in "
        "10-min increments). The target is documented as P(DCS) in percent on "
        "[0, 100].",
    )
    p = doc.add_paragraph()
    _add_inline(p, "Automated audit (`tinydcs.data_clean`) flagged two defects:")

    bullets_2_1 = [
        "**Scale inconsistency.** For 1,221 rows (7.5%), the target was entered "
        "on the fraction scale [0, 1] instead of percent. Flagging used a "
        "neighbour-median heuristic: a row's value v ≤ 1 is flagged when, within "
        "a (altitude ± 1000 ft, prebreathe ± 15 min, time-at-altitude ± 20 min, "
        "same exercise) neighbourhood of percent-scale rows, the neighbour median "
        "exceeds 1.0 and the rescaled value 100v is within ±30% of that "
        "neighbour median. Flagged rows were multiplied by 100.",
        "**Within-combo disagreement.** 26 grid cells had two rows with distinct "
        "target values; each cell was collapsed to its median.",
    ]
    for b in bullets_2_1:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        _add_inline(p, b)

    p = doc.add_paragraph()
    _add_inline(
        p,
        "After cleaning, the dataset comprised 15,908 unique grid cells. Full "
        "quality report is versioned in `artifacts/data_quality_report.md`.",
    )

    _heading(doc, "2.2 Feature vector", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "The surrogate consumes a 13-feature vector designed to be computable "
        "on-device from altitude telemetry and an accelerometer-derived VO_2 "
        "stream:",
    )
    feature_rows = [
        ["`altitude_ft`, `ambient_pressure_atm`",
         "ADRAC covariate; ISA approximation"],
        ["`prebreathe_time_min`, `prebreathe_fio2`", "Conkin / Webb [6]"],
        ["`ascent_rate_fpm`", "Depressurisation vs. slow climb"],
        ["`altitude_time_min`, `altitude_fio2`", "ADRAC covariate"],
        ["`prebreathe_vo2_mean_lmin`, `prebreathe_vo2_peak_lmin`",
         "Conkin [6] prebreathe-exercise effect"],
        ["`altitude_vo2_mean_lmin`, `altitude_vo2_peak_1min_lmin`, "
         "`altitude_vo2_integral_lmin_min`",
         "Webb [2,3] (1-min peak VO_2 in any 16-min window at altitude)"],
        ["`tissue_n2_ratio_360min`",
         "Conkin single-compartment tissue N_2 supersaturation ratio [6] "
         "(water-vapour corrected, 360-min half-time)"],
    ]
    _add_table(doc, ["Feature", "Mechanism / source"], feature_rows)

    p = doc.add_paragraph()
    _add_inline(
        p,
        "Because the ADRAC grid encodes exercise only as a categorical level, we "
        "synthesised a plausible VO_2 trajectory per row consistent with the "
        "published ranges for Rest / Mild / Heavy [2]. An Ornstein–Uhlenbeck "
        "process sampled trajectories with subject-level mean drawn from Gaussian "
        "distributions centred at 0.10 / 0.45 / 1.10 L·min⁻¹ whole-body "
        "respectively, reverting to that mean with θ = 0.3 and σ = 0.15. This "
        "synthesis is a modelling assumption documented as a reported limitation; "
        "the trajectories are deterministic given the random seed.",
    )

    _heading(doc, "2.3 Surrogate and calibration", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "The target y ∈ [0, 100] was rescaled to [0, 1] and transformed as the "
        "logit of y′ = (y(n−1) + 0.5)/n following Smithson and Verkuilen [8]. "
        "The transform prevents pathological pile-up of exact zeros, which "
        "comprise 40% of the lowest-altitude rows.",
    )
    p = doc.add_paragraph()
    _add_inline(
        p,
        "A LightGBM gradient-boosting regressor [7] fit the logit target with 400 "
        "estimators, 31 leaves, learning rate 0.05, subsample 0.9, and "
        "column-subsample 0.9. Physiological monotonicity constraints were "
        "imposed: altitude, time-at-altitude, and tissue-nitrogen-ratio features "
        "were constrained non-decreasing in the logit target; prebreathe duration "
        "and ambient pressure non-increasing.",
    )
    p = doc.add_paragraph()
    _add_inline(
        p,
        "Split-conformal prediction intervals were computed on the logit scale "
        "via the standard finite-sample quantile q = r_{(⌈(n+1)(1–α)⌉)} where "
        "α = 0.05 and residuals are computed on a held-out calibration fold [9]. "
        "To restore per-altitude-band marginal coverage under residual "
        "heteroscedasticity, we applied **Mondrian conformal** [10] stratified by "
        "5,000-ft altitude band, computing a separate quantile per band and using "
        "the overall quantile as a fallback for bands with fewer than 20 "
        "calibration samples. We also evaluated conformalised quantile regression "
        "(CQR) [11] as an ablation calibration mode.",
    )
    p = doc.add_paragraph()
    _add_inline(
        p,
        "An out-of-envelope abstention layer computes the Mahalanobis distance "
        "of a prediction input to the training-feature mean under a "
        "shrinkage-regularized covariance, flagging samples above the 99th "
        "percentile of training distances. The three-layer stack — wearable "
        "input / LightGBM logit core / conformal calibration — is schematised in "
        "Figure 1.",
    )

    _heading(doc, "2.4 Training and evaluation", level=2)
    paras_2_4 = [
        "The cleaned grid was randomly partitioned 65% / 20% / 15% into "
        "training / calibration / test folds with a fixed seed. For the "
        "closed-form ADRAC baseline, `mechanistic.adrac.fit_adrac` fit the "
        "log-logistic AFT functional form [12,1] to the target via L-BFGS-B "
        "minimisation of logit-residual mean-squared error.",
        "Evaluation reported (i) point accuracy (MAE, RMSE, R²), (ii) Brier "
        "score, (iii) calibration slope and intercept [13] (weighted logistic "
        "recalibration), (iv) Bland–Altman bias with 95% limits of agreement, "
        "(v) empirical conformal coverage overall and per altitude band, and "
        "(vi) ONNX inference latency and size on CPU as a proxy for embedded "
        "performance. Reporting follows the TRIPOD+AI guidance [16] (Appendix A).",
        "Robustness was tested under leave-one-altitude-out cross-validation: "
        "5 disjoint 5,000-ft bands, each used in turn as the held-out test set.",
    ]
    for txt in paras_2_4:
        p = doc.add_paragraph()
        _add_inline(p, txt)

    _heading(doc, "2.5 Edge deployment", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "Trained LightGBM models were exported to ONNX (target opset 13) via "
        "`onnxmltools`, dynamic-INT8 quantised via ONNX Runtime, and benchmarked "
        "on a single CPU core (AMD x86-64, but numbers in this manuscript are "
        "*indicative* — Cortex-M validation is part of ongoing work). Parity "
        "between ONNX and Python reference predictions was verified within a "
        "1e-4 max absolute error tolerance on the logit scale.",
    )

    # ----- 3. Results -----
    _heading(doc, "3. Results", level=1)

    _heading(doc, "3.1 Head-to-head on a random held-out test fold", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "On the same 2,386-row random test fold (apples-to-apples):",
    )
    _add_table(
        doc,
        ["Model", "MAE", "R²", "Brier", "Calibration slope"],
        [
            ["ADRAC closed-form AFT", "0.086", "0.869", "0.0150", "0.613"],
            ["**TinyDCS (zero-inflated)**", "**0.020**", "**0.986**",
             "**0.0016**", "**0.970**"],
        ],
        caption="**Table 1.** Head-to-head accuracy and calibration on the random "
        "held-out test fold (*n* = 2,386).",
    )
    p = doc.add_paragraph()
    _add_inline(
        p,
        "TinyDCS attains a 4-fold reduction in MAE and ~10-fold reduction in "
        "Brier score relative to the closed-form baseline, with a calibration "
        "slope (0.970) much closer to the ideal of 1.0 than the baseline's 0.613. "
        "Reliability across the full probability range is shown in Figure 2.",
    )

    _heading(doc, "3.2 Robustness to altitude extrapolation", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "Under leave-one-altitude-out cross-validation (5 bands of 5,000 ft), "
        "TinyDCS retained its advantage but the gap narrowed:",
    )
    _add_table(
        doc, ["Model", "MAE (mean ± SD)"],
        [
            ["ADRAC baseline", "0.081 ± 0.037"],
            ["**TinyDCS**", "**0.059 ± 0.033**"],
        ],
        caption="**Table 2.** Leave-one-altitude-out cross-validation MAE "
        "(mean ± SD across the five 5,000-ft bands).",
    )
    p = doc.add_paragraph()
    _add_inline(
        p,
        "A 28% reduction in MAE under strict altitude extrapolation suggests the "
        "surrogate's richer feature set generalizes to unseen altitude bands "
        "better than the four-covariate closed-form model.",
    )

    _heading(doc, "3.3 Model-size ladder and edge deployment", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "Because the zero-inflated surrogate is the production default, each "
        "deployable variant consists of a matched pair of LightGBM models "
        "(binary classifier + continuous regressor). We trained four size "
        "variants of the pair and exported both stages to ONNX.",
    )
    _add_table(
        doc,
        ["Variant", "Estim. × Leaves", "MAE", "R²", "Coverage",
         "Classifier", "Regressor", "Combined", "p50 latency"],
        [
            ["Full", "400 × 31", "0.020", "0.986", "0.960", "896 KB",
             "891 KB", "1,787 KB", "16.5 μs"],
            ["Medium", "200 × 15", "0.023", "0.984", "0.956", "212 KB",
             "211 KB", "423 KB", "6.1 μs"],
            ["**Compact**", "100 × 7", "0.028", "0.981", "0.953", "47 KB",
             "47 KB", "**95 KB**", "**2.44 μs**"],
            ["Tiny", "50 × 5", "0.033", "0.975", "0.948", "17 KB", "17 KB",
             "34 KB", "1.9 μs"],
        ],
        caption="**Table 3.** TinyDCS size-ladder variants: ONNX size (per stage "
        "and combined) and CPU per-row inference latency.",
    )
    p = doc.add_paragraph()
    _add_inline(
        p,
        "The **Compact** variant is the headline: it achieves the 100 KB combined "
        "ONNX footprint targeted for smartwatch and flight-watch deployment "
        "while retaining R² > 0.98, closing the low-band coverage shortfall "
        "(see §3.5), and still outperforming the closed-form ADRAC baseline by "
        "3× on MAE. The Pareto frontier of ONNX size versus MAE is shown in "
        "Figure 3.",
    )

    _heading(doc, "3.4 Inference latency and embedded feasibility", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "CPU per-row inference latency (single AMD x86-64 core, batch 10,000 "
        "rows) for the Compact zero-inflated pair:",
    )
    for b in ["p50 = 2.44 μs / row", "p95 = 3.31 μs / row"]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.add_run(b)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "The Full variant measured 16.5 μs p50. ONNX-to-Python parity was "
        "verified within 6.4 × 10⁻⁷ max absolute error on P(y = 0) and "
        "4.9 × 10⁻⁶ on the continuous logit — well below the 10⁻⁴ target. "
        "Per-row latency on a Cortex-M4 at 80 MHz is expected to be ~20× slower "
        "than server CPU under TFLite Micro-style quantised kernels [14], giving "
        "the Compact variant an operational budget of ~50 μs / inference on "
        "bare metal — two orders of magnitude below the 1 ms wearable-alerting "
        "target. Direct Cortex-M validation remains part of ongoing work.",
    )

    _heading(doc, "3.5 Calibration", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "We compared five calibration strategies on the same random test fold. "
        "Per-altitude-band coverage is the diagnostic.",
    )
    _add_table(
        doc,
        ["Calibration", "Overall", "18–23K ft", "23–28K ft", "28–33K ft",
         "33–38K ft", "38–43K ft"],
        [
            ["Global conformal", "0.869", "0.591", "0.933", "0.944", "0.967",
             "0.948"],
            ["Mondrian", "0.869", "0.583", "0.949", "0.945", "0.954", "0.955"],
            ["CQR (global q)", "0.864", "0.589", "0.937", "0.945", "0.945",
             "0.937"],
            ["Mondrian-CQR", "0.865", "0.589", "0.924", "0.959", "0.951",
             "0.937"],
            ["**Zero-inflated two-stage**", "**0.960**", "**0.964**",
             "**0.953**", "**0.951**", "**0.967**", "**0.966**"],
        ],
        caption="**Table 4.** Per-altitude-band 95% conformal coverage across "
        "five calibration strategies on the held-out test fold.",
    )
    p = doc.add_paragraph()
    _add_inline(
        p,
        "Four conformal-only methods produce near-nominal coverage in the four "
        "upper bands (0.92–0.97) but are invariant at 0.58–0.59 in the lowest "
        "altitude band. The invariance across methods is diagnostic: the "
        "shortfall is target-distribution pathology rather than residual "
        "variance. Routing the ~40% exact-zero mass through a dedicated binary "
        "classifier (stage 1), with the continuous regressor (stage 2) trained "
        "only on non-zero rows, closes the gap entirely — the "
        "18,000–23,000 ft band reaches 0.964 and overall coverage is 0.960. "
        "The zero-inflated model is therefore adopted as the default calibration "
        "for all headline numbers; the three conformal-only modes remain "
        "available as ablations. Per-altitude-band coverage profiles for all "
        "five strategies are shown in Figure 4.",
    )

    _heading(doc, "3.6 Personalisation prototype", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "A conjugate-Gaussian hierarchical model placed a Gaussian prior on "
        "per-subject log-susceptibility (prior mean and variance fit to the "
        "population distribution) and updated it via Bayes' rule after each "
        "observed exposure. On a synthetic 200-subject cohort (true "
        "susceptibilities drawn from the prior, DCS outcomes simulated with the "
        "TinyDCS point estimate as the population rate), the model recovered "
        "per-subject log-susceptibility at Pearson *r* = 0.63 after twenty "
        "exposures per subject. Population-level versus personalized Brier "
        "parity crossover occurred near *k* = 10 exposures. These results "
        "establish feasibility and scope for Paper 2 (full hierarchical "
        "Bayesian model on real chamber cohort data). The information-gain "
        "curve is shown in Figure 5.",
    )

    _heading(doc, "3.7 Dataset quality", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "The cleaner repaired 1,221 rows (7.5% of the raw dataset) with a "
        "systematic fraction-vs-percent scale inconsistency. Before cleaning, "
        "the target variable had max 98 but 2,348 rows with values in (0, 1], "
        "indicating the mixed-scale pathology. After cleaning, all targets were "
        "in the expected [0, 100] range with sensible gradients across the "
        "factorial grid. 26 within-combo disagreements were resolved to the "
        "median. Full report versioned as `artifacts/data_quality_report.md`.",
    )

    # ----- 4. Discussion -----
    _heading(doc, "4. Discussion", level=1)

    _heading(doc, "4.1 What the numbers mean", level=2)
    paras_4_1 = [
        "A MAE of 0.020 on held-out ADRAC-output probability is approximately "
        "one order of magnitude below the typical resolution of operational "
        "aeromedical decisions (which are made in coarse increments of perhaps "
        "5%). The R² of 0.986 on a factorial grid reflects that the surrogate "
        "is doing what it should — recovering a known function — but the "
        "leave-one-altitude-out MAE of 0.059 is the more honest measure of "
        "generalization, and even there TinyDCS reduces the ADRAC baseline's "
        "error by 28%.",
        "A 95 KB combined ONNX footprint with 2.44 μs p50 inference fits "
        "comfortably on every major wearable platform — a smartwatch main "
        "processor executes more than a million such inferences per second "
        "while still meeting its display and notification duties. The "
        "calibrated zero-inflated interval plus Mahalanobis envelope abstention "
        "provides the kind of uncertainty-aware output that a pilot or chamber "
        "medic can reason about operationally: a predicted 12% DCS risk "
        "accompanied by “[3%, 28%] 95% CI, in-envelope” is actionable in a "
        "way that a point estimate alone is not.",
    ]
    for t in paras_4_1:
        p = doc.add_paragraph()
        _add_inline(p, t)

    _heading(doc, "4.2 Why the surrogate beats the closed-form baseline", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "The closed-form ADRAC baseline is a four-covariate log-logistic AFT "
        "fit to a factorial grid. Its functional form bakes in a specific "
        "interaction structure between log-time and covariates. The LightGBM "
        "surrogate relaxes that structure while remaining physiologically "
        "well-behaved via monotonicity constraints, and its 13-feature input "
        "captures mechanism (tissue nitrogen ratio) and exposure detail "
        "(continuous VO_2 trajectory summaries) that the closed-form model "
        "cannot. The improvement is expected and, importantly, does not come "
        "at the cost of deployment complexity: the combined ONNX pair is under "
        "100 KB.",
    )

    _heading(doc, "4.3 Limitations", level=2)
    paras_4_3 = [
        "*Ground truth is a mechanistic model's output, not observed DCS "
        "incidence.* TinyDCS is a surrogate of ADRAC; any claim beyond "
        "“reproduces ADRAC with richer covariates and calibrated "
        "uncertainty” requires prospective validation against chamber-observed "
        "symptoms. That work is the subject of a planned external-validation "
        "study in the Colombian CEMAE hypobaric chamber with Doppler venous-gas-"
        "emboli ground truth (Paper 3 scope).",
        "*Zero-inflated calibration was necessary; mean-only calibration would "
        "not have sufficed.* Early variants produced an invariant 0.583 coverage "
        "in the 18,000–23,000 ft band across four conformal-only calibration "
        "modes (global, Mondrian, CQR, Mondrian-CQR). The invariance of the "
        "shortfall across methods was diagnostic: it reflected target-distribution "
        "pathology, not residual-variance error. Because ~40% of rows in the "
        "low-altitude band have target exactly zero, a mean regressor trained on "
        "the whole grid cannot represent that mass cleanly. The two-stage "
        "zero-inflated architecture — a LightGBM classifier for "
        "P(y = 0 | x) gating a continuous regressor trained only on non-zero "
        "rows — closes the gap entirely and is adopted as the production "
        "default. The three conformal-only calibration modes remain available "
        "in the public API as ablations; Mondrian-CQR in particular produces "
        "intervals ~50% narrower than Mondrian conformal in the four upper "
        "altitude bands while maintaining coverage, which is a genuine "
        "efficiency advantage when a calibrated-but-narrow interval is "
        "preferred to the zero-inflated default.",
        "*Continuous-VO_2 trajectories are synthesised, not measured.* Because "
        "the ADRAC grid encodes exercise categorically, this work's "
        "continuous-VO_2 features are plausibility-grounded synthesis "
        "consistent with Webb 2010 ranges, not real wearable data. The "
        "inference-time API accepts real accelerometer-derived VO_2; external "
        "validation in a wearable-instrumented chamber cohort is required to "
        "demonstrate operational accuracy under real VO_2 signal.",
        "*Fixed validity envelope.* The surrogate abstains outside the training "
        "envelope (altitude 18,000–40,000 ft; prebreathe 0–180 min; "
        "time-at-altitude 10–240 min). It does not extrapolate.",
        "*No individual variability.* Neither TinyDCS nor the ADRAC baseline "
        "represents inter-subject differences in DCS susceptibility. "
        "Hierarchical Bayesian personalisation from multimodal wearable "
        "telemetry is the subject of a companion manuscript in preparation.",
    ]
    for t in paras_4_3:
        p = doc.add_paragraph()
        _add_inline(p, t)

    _heading(doc, "4.4 Operational implications", level=2)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "For unpressurised general-aviation operations above FL180 — a "
        "population with documented but under-monitored DCS risk [15] — "
        "TinyDCS offers a practical on-board monitor: a smartwatch or flight "
        "computer can compute altitude-informed, activity-aware risk estimates "
        "in real time, log the exposure, and alert before a predicted risk "
        "threshold is crossed. The model's refusal to predict outside its "
        "validated envelope is as important as its predictions within it — a "
        "flight exceeding the model's tested altitudes will be flagged rather "
        "than extrapolated against.",
    )

    # ----- 5. Conclusion -----
    _heading(doc, "5. Conclusion", level=1)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "We have developed a 95 KB combined, 2.44 μs/row edge-deployable "
        "machine-learning surrogate of the US Air Force ADRAC altitude-DCS risk "
        "model with continuous-VO_2 covariates, zero-inflated conformal "
        "prediction intervals with uniform altitude-band coverage, and "
        "principled Mahalanobis-distance out-of-envelope abstention. The "
        "surrogate reduces mean absolute error by 4× and Brier score by 10× "
        "relative to a closed-form log-logistic baseline on the same data, "
        "while preserving a calibration slope close to 1.0. A "
        "conjugate-Gaussian hierarchical personalisation prototype quantifies "
        "the per-subject exposure count at which individualisation matches the "
        "population model (k ≈ 10), setting the scope for a companion "
        "manuscript. The work is a prerequisite for wearable DCS monitoring in "
        "unpressurised aviation, chamber training, and extravehicular activity; "
        "prospective external validation in a Colombian chamber cohort is the "
        "immediate next step.",
    )

    # ----- Author contributions / declarations -----
    _heading(doc, "Author contributions (CRediT)", level=1)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "**Diego Malpica:** Conceptualisation, Methodology, Software, "
        "Validation, Formal analysis, Investigation, Data curation, Writing — "
        "original draft, Writing — review and editing, Visualisation, "
        "Supervision, Project administration.",
    )
    p = doc.add_paragraph()
    _add_inline(
        p,
        "**Marian Farfán:** Conceptualisation, Methodology, Investigation, "
        "Writing — review and editing, Resources.",
    )
    p = doc.add_paragraph()
    _add_inline(p, "All authors read and approved the final version of the manuscript.")

    _heading(doc, "Use of generative AI", level=1)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "During the preparation of this work, the authors used Claude Code "
        "(Anthropic Claude Opus 4.7) as a coding and prose-revision assistant "
        "under direct human supervision. AI tools were used to (i) accelerate "
        "boilerplate code (e.g. ONNX export wrappers and plotting scripts), "
        "(ii) suggest manuscript phrasing on non-scientific passages, and "
        "(iii) cross-check formatting against journal guidelines. AI tools were "
        "not used to generate scientific content, design experiments, fabricate "
        "or alter data, draft results or interpretation, or produce figures, "
        "tables, or citations without manual verification. All scientific claims "
        "and numerical results were derived from the released code base and "
        "verified by the authors. After using these tools, the authors reviewed "
        "and edited the content as needed and take full responsibility for the "
        "content of the publication.",
    )

    _heading(doc, "Declaration of competing interests", level=1)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "The authors declare that they have no known competing financial "
        "interests or personal relationships that could have appeared to "
        "influence the work reported in this paper.",
    )

    _heading(doc, "Funding", level=1)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "This work received no external funding. The work was conducted as "
        "part of the authors' duties at the Colombian Aerospace Force, "
        "Direction of Aerospace Medicine.",
    )

    _heading(doc, "Data and code availability", level=1)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "All code, the cleaned ADRAC-derived dataset, trained model bundles "
        "(joblib), ONNX artefacts at four size tiers, metrics JSONs, paper "
        "figures, the TRIPOD+AI checklist, and a command-by-command reproduction "
        "guide are released under a research-use license at "
        "https://github.com/strikerdlm/DCS (default branch `main`, tagged "
        "release at the time of submission). All training runs are reproducible "
        "from the shipped raw CSV with `seed = 42` in under three minutes on "
        "CPU via `docs/runbook.md`. No restricted data are involved.",
    )

    _heading(doc, "Acknowledgements", level=1)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "The authors thank the Jefatura de Educación Aeronáutica y Espacial "
        "(Aeronautics and Space Education Directorate) and the Jefatura de "
        "Salud (Health Directorate) of the Colombian Aerospace Force for their "
        "institutional support of this research. We are grateful to CR Julio "
        "Blanco, MY Nindre Pico, MY Luis Eduardo Jerez, TS Bernabé Cardona, "
        "T1 Ever Buitrago, T1 Angie Alvarado, and T1 Rafael Salamanca for their "
        "thoughtful review, technical input, and the team effort that made this "
        "work possible.",
    )

    # ----- References -----
    _heading(doc, "References", level=1)
    references = [
        "Pilmanis AA, Petropoulos L, Kannan N, Webb JT. Decompression sickness "
        "risk model: development and validation by 150 prospective hypobaric "
        "exposures. Aviat Space Environ Med 2004;75(9):749–59. PMID: 15460625.",
        "Webb JT, Krock LP, Gernhardt ML. Oxygen consumption at altitude as a "
        "risk factor for altitude decompression sickness. Aviat Space Environ "
        "Med 2010;81(11):987–92. https://doi.org/10.3357/ASEM.2787.2010.",
        "Webb JT, Morgan TR, Sarsfield SD. Altitude decompression sickness "
        "risk and physical activity during exposure. Aerosp Med Hum Perform "
        "2016;87:516–20. https://doi.org/10.3357/AMHP.4477.2016.",
        "Gerth WA, Doolette DJ, Gault KA. A probabilistic model of altitude "
        "decompression sickness based on the 3RUT-MB model. Navy Experimental "
        "Diving Unit Technical Report NEDU TR 18-01; 2018. DTIC accession "
        "AD1101527.",
        "Han Y, et al. Machine learning methods to predict incidence risk of "
        "altitude decompression sickness. Proc IEEE Int Conf Comput Aerosp "
        "Robot Eng (CACRE); 2023.",
        "Conkin J, Gernhardt ML. A probability model of decompression sickness "
        "at 4.3 psia after exercise prebreathe. NASA Technical Paper "
        "TP-2004-213158; 2004.",
        "Ke G, Meng Q, Finley T, Wang T, Chen W, Ma W, et al. LightGBM: a "
        "highly efficient gradient boosting decision tree. Adv Neural Inf "
        "Process Syst (NeurIPS) 2017;30:3146–54.",
        "Smithson M, Verkuilen J. A better lemon squeezer? Maximum-likelihood "
        "regression with beta-distributed dependent variables. Psychol Methods "
        "2006;11(1):54–71. https://doi.org/10.1037/1082-989X.11.1.54.",
        "Shafer G, Vovk V. A tutorial on conformal prediction. J Mach Learn "
        "Res 2008;9:371–421.",
        "Vovk V, Gammerman A, Shafer G. Algorithmic learning in a random "
        "world. 2nd ed. Springer; 2022 (Mondrian conformal, Chapter 4). "
        "https://doi.org/10.1007/978-3-031-06649-8.",
        "Romano Y, Patterson E, Candès EJ. Conformalized quantile regression. "
        "Adv Neural Inf Process Syst (NeurIPS) 2019;32:3543–53.",
        "Kannan N, Raychaudhuri A, Pilmanis AA. A loglogistic model for "
        "altitude decompression sickness. Aviat Space Environ Med "
        "1998;69:965–70.",
        "Van Calster B, McLernon DJ, van Smeden M, Wynants L, Steyerberg EW; "
        "Topic Group ‘Evaluating diagnostic tests and prediction models’ of "
        "the STRATOS initiative. Calibration: the Achilles heel of predictive "
        "analytics. BMC Med 2019;17:230. "
        "https://doi.org/10.1186/s12916-019-1466-7.",
        "Warden P, Situnayake D. TinyML: machine learning with TensorFlow Lite "
        "on Arduino and ultra-low-power microcontrollers. Sebastopol (CA): "
        "O’Reilly Media; 2019.",
        "Harrison MF, Butler WP, Stepanek J. Decompression sickness risk "
        "assessment and awareness in general aviation. Aerosp Med Hum Perform "
        "2021;92(3):138–45.",
        "Collins GS, Moons KGM, Dhiman P, Riley RD, Beam AL, Van Calster B, "
        "et al. TRIPOD+AI statement: updated guidance for reporting clinical "
        "prediction models that use regression or machine learning methods. "
        "BMJ 2024;385:e078378. https://doi.org/10.1136/bmj-2023-078378.",
    ]
    for i, ref in enumerate(references, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"[{i}] ")
        run.bold = True
        _add_inline(p, ref)

    # ----- Appendix A — TRIPOD+AI -----
    doc.add_page_break()
    _heading(doc, "Appendix A — TRIPOD+AI [16] checklist coverage", level=1)
    tripod_rows = [
        ["1. Title identifies model type and context", "Title page"],
        ["2. Abstract structured", "Abstract"],
        ["3. Background + rationale", "§1"],
        ["4. Objectives", "§1 Contributions"],
        ["5. Data source and eligibility", "§2.1"],
        ["6. Outcome definition", "§2.1 (ADRAC model output)"],
        ["7. Predictors", "§2.2"],
        ["8. Sample size", "§2.1 (15,908 after cleaning)"],
        ["9. Missing data handling", "§2.1 (cleaner + dedup)"],
        ["10. Statistical / ML methods", "§2.3"],
        ["11. Training, tuning, validation", "§2.4"],
        ["12. Performance measures", "§2.4"],
        ["13. Model specification", "§2.3 + released artifacts"],
        ["14. Performance results", "§3"],
        ["15. Calibration, discrimination, clinical utility",
         "§3.1–3.5, Fig. 1–3"],
        ["16. Limitations", "§4.3"],
        ["17. Interpretation and generalisability", "§4"],
        ["18. Data / code availability", "Data and code availability"],
        ["19. Funding and conflicts",
         "Funding; Declaration of competing interests"],
    ]
    _add_table(doc, ["Item", "Location"], tripod_rows)

    # ----- Figure captions -----
    _heading(doc, "Figure captions", level=1)
    captions = [
        ("Figure 1.",
         "TinyDCS system architecture. Block diagram of the three-layer "
         "inference stack: (1) wearable sensor input layer producing a "
         "13-feature vector from altitude telemetry and accelerometer-derived "
         "VO_2; (2) LightGBM logit core with monotonicity constraints and "
         "Mahalanobis OOD gate; (3) zero-inflated conformal calibration layer "
         "returning a point estimate plus a calibrated 95% interval. ONNX "
         "artefacts are shown at the edge-deployment node."),
        ("Figure 2.",
         "Reliability diagram — TinyDCS vs. closed-form ADRAC baseline. "
         "Predicted P(DCS) bins (x-axis) versus empirical observed fraction "
         "(y-axis) on the held-out test fold (*n* = 2,386). Perfect "
         "calibration lies on the diagonal. TinyDCS (zero-inflated two-stage) "
         "tracks the diagonal closely across the full probability range; the "
         "closed-form AFT baseline shows systematic overestimation at low "
         "probabilities."),
        ("Figure 3.",
         "ONNX model size versus MAE — Pareto frontier across the size "
         "ladder. Log-scale x-axis (ONNX file size in KB) versus MAE on the "
         "held-out test fold. Four TinyDCS variants (Tiny, Compact, Medium, "
         "Full) and the closed-form ADRAC baseline are plotted. The Compact "
         "variant achieves the target edge-deployment footprint while "
         "dominating the baseline by 3× on MAE."),
        ("Figure 4.",
         "Per-altitude-band 95% conformal coverage — five calibration "
         "strategies. Grouped bars showing empirical coverage in each "
         "5,000-ft altitude band for five calibration methods on the same "
         "test fold. Four conformal-only methods (global, Mondrian, CQR, "
         "Mondrian-CQR) are invariant at 0.58–0.59 in the 18,000–23,000 ft "
         "band. The zero-inflated two-stage method achieves ≥ 0.95 coverage "
         "in all five bands."),
        ("Figure 5.",
         "Personalisation information gain — per-subject susceptibility "
         "recovery. Left y-axis: Pearson *r* between true and posterior-mean "
         "log-susceptibility (synthetic 200-subject cohort) as a function of "
         "observed exposures *k* per subject. Right y-axis: Brier score for "
         "population-level (flat prior) versus personalised predictions. "
         "Crossover near *k* = 10 indicates the exposure count at which "
         "personalisation begins to outperform the population model."),
    ]
    for label, body in captions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(label + " ")
        run.bold = True
        _add_inline(p, body)

    out = OUT / "Manuscript.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ---------------------------------------------------------------------------
# Cover letter
# ---------------------------------------------------------------------------

def build_cover_letter() -> None:
    doc = _new_doc(double_spaced=False, line_numbers=False)

    # Header (institutional)
    _add_para(doc,
              "Subdirectorate of Aerospace Sciences",
              align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, space_after=0)
    _add_para(doc,
              "Direction of Aerospace Medicine",
              align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _add_para(doc,
              "Colombian Aerospace Force",
              align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _add_para(doc,
              "Bogotá DC, Colombia",
              align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

    _add_para(doc, "2 May 2026", space_after=18)

    _add_para(doc, EIC, bold=True, space_after=0)
    for line in EIC_AFFIL.split("\n"):
        _add_para(doc, line, space_after=0)
    _add_para(doc, "", space_after=12)
    _add_para(doc, "Submitted via Editorial Manager "
              "(https://www.editorialmanager.com/cmpb/)",
              italic=True, space_after=18)

    _add_para(doc, "Dear Prof. Molinari,", space_after=12)

    intro = doc.add_paragraph()
    _add_inline(
        intro,
        "We are pleased to submit our manuscript, *“TinyDCS: An "
        "edge-deployable machine-learning surrogate of the ADRAC "
        "altitude-decompression-sickness risk model with continuous-exposure "
        "covariates and calibrated uncertainty,”* for consideration as a "
        "**Full Length Article** in *Computer Methods and Programs in "
        "Biomedicine*.",
    )
    intro.paragraph_format.space_after = Pt(12)

    sections = [
        ("Scope and fit.",
         "The manuscript is, at its core, a methodological and software "
         "contribution: a monotonicity-constrained gradient-boosting surrogate "
         "of a published survival-model probability grid, equipped with a "
         "two-stage zero-inflated split-conformal calibration, Mondrian "
         "altitude-band stratification, Mahalanobis-distance "
         "out-of-distribution abstention, and an ONNX export pipeline that "
         "compiles to 95 KB and runs at 2.44 μs per inference on commodity "
         "CPU. Altitude decompression sickness in unpressurised aviation and "
         "extravehicular activity is the demonstrating use case. The "
         "methodology — boundary-shrinkage logit regression with "
         "monotonicity constraints, Mondrian conformal calibration with a "
         "dedicated zero-inflation stage to repair coverage where the target "
         "distribution is degenerate, and OOD-aware abstention for "
         "safety-critical inference — generalises to any biomedical "
         "risk-grid surrogate with similar pathologies. CMPB’s mandate to "
         "encourage formal computing methods and their application in "
         "biomedical research and medical practice is the natural home for "
         "this work."),
        ("What is known.",
         "The US Air Force ADRAC log-logistic accelerated-failure-time model "
         "(Pilmanis 2004) is the operational standard for altitude DCS risk "
         "planning, but it accepts only a three-level categorical exercise "
         "covariate and emits point estimates without calibrated uncertainty. "
         "The bubble-dynamics 3RUT-MBe1 model (NEDU TR 18-01, 2018) accepts "
         "continuous VO_2 trajectories but is too computationally heavy for "
         "embedded deployment."),
        ("What this study adds.",
         "TinyDCS is, to our knowledge, the first published model that "
         "simultaneously (i) reproduces the ADRAC grid with a 4-fold MAE "
         "reduction (0.020 vs 0.086) and 10-fold Brier reduction over a "
         "closed-form log-logistic AFT fit to the same data, (ii) accepts "
         "continuous wearable-derived VO_2 covariates, (iii) ships calibrated "
         "95% prediction intervals with uniform per-altitude-band coverage "
         "(closing a 0.58 → 0.96 shortfall in the 18–23 kft band that four "
         "conformal-only alternatives could not close), (iv) abstains "
         "principled outside its validated input envelope, and (v) compiles "
         "to a 95 KB ONNX artefact with 2.44 μs per-row inference latency. "
         "A reproducibility package (runbook, AI-agent continuation guide, "
         "trained weights, ONNX exports, metrics JSONs, TRIPOD+AI checklist) "
         "is publicly released at https://github.com/strikerdlm/DCS."),
        ("Generalisability.",
         "The two-stage zero-inflated conformal architecture and the "
         "Mahalanobis envelope abstention are agnostic to the aerospace "
         "domain and applicable to any survival-model surrogate with a "
         "degenerate target mass at one boundary — a common pathology in "
         "clinical risk modelling."),
    ]
    for title, body in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(title + " ")
        run.bold = True
        _add_inline(p, body)

    _add_para(doc, "Declarations.", bold=True, space_after=4)

    decls = [
        ("Originality.",
         "The manuscript has not been previously published and is not under "
         "concurrent consideration at any other journal."),
        ("AI disclosure.",
         "Generative AI tools (Claude Code, Anthropic) were used as a coding "
         "and prose-revision assistant under direct human supervision. AI did "
         "not generate scientific content, did not author the manuscript, and "
         "was not used to fabricate data, citations, figures, or analyses. "
         "All scientific claims and numerical results were derived from the "
         "released code base and verified by the authors. No AI-generated "
         "text was retained without manual review."),
        ("Conflicts of interest.",
         "All authors declare no conflicts of interest."),
        ("Funding.",
         "This work received no external funding."),
        ("Ethical approval.",
         "Not applicable. The study uses a published computational grid "
         "(ADRAC) and synthetic VO_2 trajectories; no human or animal "
         "subjects were involved."),
        ("Data and code availability.",
         "Repository: https://github.com/strikerdlm/DCS. License: "
         "research-use. Reproduction is end-to-end deterministic via "
         "docs/runbook.md (~3 minutes on CPU)."),
        ("Suggested reviewers.",
         "Five candidates with relevant expertise in conformal prediction in "
         "clinical contexts, ML surrogates of physiological/biomedical risk "
         "models, and TinyML for wearable physiological monitoring are listed "
         "in the accompanying suggested-reviewers file. None has co-authored "
         "with us in the past three years; none shares our institutional "
         "affiliation; none is a member of the CMPB editorial board to our "
         "knowledge."),
        ("Software product.",
         "The TinyDCS Python library, its ONNX exports at four size tiers "
         "(17 KB / 47 KB / 211 KB / 891 KB regressor + paired classifier), "
         "and a `tinydcs.zi` API for the production zero-inflated configuration "
         "are released open-source as part of this submission. The release "
         "supports CPU and is positioned for Cortex-M validation in follow-up "
         "work."),
    ]
    for title, body in decls:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title + " ")
        run.bold = True
        _add_inline(p, body)

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    _add_inline(
        closing,
        "We confirm that all authors have read and approved the submitted "
        "version, and that the manuscript reports our original work. We thank "
        "you and the reviewers for your time and consideration.",
    )

    _add_para(doc, "", space_after=12)
    _add_para(doc, "Sincerely,", space_after=18)

    _add_para(doc, "Diego Malpica, MD", bold=True, space_after=0)
    _add_para(doc, "Corresponding author", italic=True, space_after=0)
    _add_para(doc, "Subdirectorate of Aerospace Sciences", space_after=0)
    _add_para(doc, "Direction of Aerospace Medicine, Colombian Aerospace Force",
              space_after=0)
    _add_para(doc, "Bogotá DC, Colombia", space_after=0)
    _add_para(doc, "Email: diego.malpica@fac.mil.co", space_after=12)

    _add_para(doc, "Marian Farfán, MD", bold=True, space_after=0)
    _add_para(doc, "Subdirectorate of Aerospace Sciences", space_after=0)
    _add_para(doc, "Direction of Aerospace Medicine, Colombian Aerospace Force",
              space_after=0)
    _add_para(doc, "Bogotá DC, Colombia", space_after=0)

    out = OUT / "Cover Letter.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ---------------------------------------------------------------------------
# Highlights
# ---------------------------------------------------------------------------

def build_highlights() -> None:
    doc = _new_doc(double_spaced=False, line_numbers=False)

    _add_para(doc, "Highlights",
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16,
              space_after=4)
    _add_para(doc,
              "TinyDCS: An edge-deployable machine-learning surrogate of the "
              "ADRAC altitude-decompression-sickness risk model with "
              "continuous-exposure covariates and calibrated uncertainty",
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    bullets = [
        "TinyDCS surrogate of ADRAC reaches MAE 0.020 vs 0.086 for closed-form baseline.",
        "Compact 95 KB ONNX model with 2.44 μs/row CPU latency for wearable deployment.",
        "Zero-inflated two-stage calibration restores 95% coverage at low altitudes.",
        "Continuous-VO_2 covariates with monotonicity constraints and OOD abstention.",
        "Bayesian personalisation prototype reaches r = 0.63 after 20 per-subject runs.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(6)
        _add_inline(p, b)

    out = OUT / "Highlights.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ---------------------------------------------------------------------------
# Declaration of Competing Interest (standalone Elsevier form)
# ---------------------------------------------------------------------------

def build_coi() -> None:
    doc = _new_doc(double_spaced=False, line_numbers=False)

    _add_para(doc, "Declaration of Competing Interest",
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16,
              space_after=4)
    _add_para(doc, JOURNAL,
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    _add_para(doc, "Manuscript title:", bold=True, space_after=0)
    p = doc.add_paragraph()
    _add_inline(p, TITLE)
    p.paragraph_format.space_after = Pt(12)

    _add_para(doc, "Authors:", bold=True, space_after=0)
    _add_para(doc, "Diego Malpica, MD; Marian Farfán, MD",
              space_after=18)

    p = doc.add_paragraph()
    _add_inline(
        p,
        "The authors whose names are listed immediately below certify that "
        "they have **no affiliations with or involvement in any organisation "
        "or entity with any financial interest or non-financial interest "
        "(such as personal or professional relationships, affiliations, "
        "knowledge or beliefs)** in the subject matter or materials discussed "
        "in this manuscript.",
    )
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    _add_inline(
        p,
        "The authors declare that they have no known competing financial "
        "interests or personal relationships that could have appeared to "
        "influence the work reported in this paper.",
    )
    p.paragraph_format.space_after = Pt(18)

    _add_para(doc, "Author signatures.", bold=True, space_after=12)

    for name in ("Diego Malpica, MD (corresponding author)",
                 "Marian Farfán, MD"):
        _add_para(doc, f"Name: {name}", space_after=4)
        _add_para(doc, "Affiliation: Subdirectorate of Aerospace Sciences, "
                  "Direction of Aerospace Medicine, Colombian Aerospace "
                  "Force, Bogotá DC, Colombia", space_after=4)
        _add_para(doc, "Signature: ____________________________________     "
                  "Date: 2 May 2026", space_after=18)

    out = OUT / "Declaration of Competing Interest.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ---------------------------------------------------------------------------
# Statement on Studies in Humans and Animals
# ---------------------------------------------------------------------------

def build_human_animal() -> None:
    doc = _new_doc(double_spaced=False, line_numbers=False)

    _add_para(doc, "Statement on Studies Involving Human and Animal Subjects",
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16,
              space_after=4)
    _add_para(doc, JOURNAL,
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    _add_para(doc, "Manuscript title:", bold=True, space_after=0)
    p = doc.add_paragraph()
    _add_inline(p, TITLE)
    p.paragraph_format.space_after = Pt(12)

    _add_para(doc, "Authors:", bold=True, space_after=0)
    _add_para(doc, "Diego Malpica, MD; Marian Farfán, MD", space_after=18)

    _add_para(doc, "1. Studies involving human participants",
              bold=True, space_after=4)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "**Not applicable.** This study did not involve any human "
        "participants, identifiable human material, or identifiable human "
        "data. No prospective or retrospective recruitment, intervention, "
        "questionnaire, interview, biological sampling, or chart review of "
        "human subjects was performed. The work uses (i) the publicly "
        "released ADRAC computational grid (`DCS_Risk_DB_2025.csv`), which "
        "represents the deterministic output of the U.S. Air Force "
        "altitude-decompression-sickness risk model on a factorial grid of "
        "exposures, and (ii) synthetic VO_2 trajectories generated by an "
        "Ornstein–Uhlenbeck process consistent with published activity-level "
        "ranges. No human-subjects approval was therefore required. The "
        "Declaration of Helsinki and applicable national regulations on "
        "research involving human participants do not apply to this work.",
    )
    p.paragraph_format.space_after = Pt(12)

    _add_para(doc, "2. Studies involving animals", bold=True, space_after=4)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "**Not applicable.** No animal experiments, animal-derived tissues, "
        "or animal-derived data were used in this study. ARRIVE guidelines "
        "and applicable national regulations on the protection of animals "
        "used for scientific purposes do not apply to this work.",
    )
    p.paragraph_format.space_after = Pt(12)

    _add_para(doc, "3. Informed consent", bold=True, space_after=4)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "**Not applicable.** Because no human participants or identifiable "
        "human data were involved, no informed-consent procedure was "
        "required.",
    )
    p.paragraph_format.space_after = Pt(12)

    _add_para(doc, "4. Ethics committee / IRB approval",
              bold=True, space_after=4)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "**Not applicable.** Approval from an Institutional Review Board, "
        "Ethics Committee, or Animal Care Committee was not required for "
        "this purely computational study.",
    )
    p.paragraph_format.space_after = Pt(18)

    _add_para(doc, "Author confirmation.", bold=True, space_after=12)
    p = doc.add_paragraph()
    _add_inline(
        p,
        "The authors confirm that the above statements are accurate and that "
        "this manuscript is a strictly computational study that does not "
        "involve human or animal subjects in any phase of its design, "
        "conduct, or analysis.",
    )
    p.paragraph_format.space_after = Pt(18)

    for name in ("Diego Malpica, MD (corresponding author)",
                 "Marian Farfán, MD"):
        _add_para(doc, f"Name: {name}", space_after=4)
        _add_para(doc, "Signature: ____________________________________     "
                  "Date: 2 May 2026", space_after=12)

    out = OUT / "Statement on Human and Animal Studies.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ---------------------------------------------------------------------------
# Author Contributions (CRediT) standalone file
# ---------------------------------------------------------------------------

def build_credit() -> None:
    doc = _new_doc(double_spaced=False, line_numbers=False)

    _add_para(doc, "Author Contributions (CRediT)",
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16,
              space_after=4)
    _add_para(doc, JOURNAL,
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    _add_para(doc, "Manuscript title:", bold=True, space_after=0)
    p = doc.add_paragraph()
    _add_inline(p, TITLE)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    _add_inline(
        p,
        "Author contributions are reported using the CRediT (Contributor "
        "Roles Taxonomy, https://credit.niso.org) following the recommendation "
        "of the National Information Standards Organization.",
    )
    p.paragraph_format.space_after = Pt(12)

    _add_table(
        doc, ["Author", "CRediT roles"],
        [
            ["Diego Malpica, MD",
             "Conceptualisation; Methodology; Software; Validation; Formal "
             "analysis; Investigation; Data curation; Writing — original "
             "draft; Writing — review and editing; Visualisation; "
             "Supervision; Project administration."],
            ["Marian Farfán, MD",
             "Conceptualisation; Methodology; Investigation; Writing — "
             "review and editing; Resources."],
        ],
    )

    p = doc.add_paragraph()
    _add_inline(p, "All authors read and approved the final version of the manuscript.")
    p.paragraph_format.space_after = Pt(18)

    out = OUT / "Author Contributions (CRediT).docx"
    doc.save(out)
    print(f"Wrote {out}")


# ---------------------------------------------------------------------------
# Suggested Reviewers (clean docx for portal entry)
# ---------------------------------------------------------------------------

def build_reviewers() -> None:
    doc = _new_doc(double_spaced=False, line_numbers=False)

    _add_para(doc, "Suggested Reviewers",
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16,
              space_after=4)
    _add_para(doc, JOURNAL,
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    _add_para(doc, "Manuscript title:", bold=True, space_after=0)
    p = doc.add_paragraph()
    _add_inline(p, TITLE)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    _add_inline(
        p,
        "Five suggested reviewers are listed below. None has co-authored "
        "with the authors in the past three years, none shares the authors’ "
        "institutional affiliation, and none is a member of the CMPB "
        "editorial board to the authors’ knowledge. All institutional "
        "emails were verified at primary sources on 2026-05-01.",
    )
    p.paragraph_format.space_after = Pt(12)

    reviewers = [
        {
            "n": 1,
            "name": "Peter H. Charlton, PhD",
            "affil": "Senior Research Scientist, Nokia Bell Labs Cambridge, UK; "
                     "previously British Heart Foundation Research Fellow, "
                     "Department of Public Health and Primary Care, "
                     "University of Cambridge",
            "email": "pete@oxon.org",
            "orcid": "0000-0003-3836-8655",
            "rationale":
                "Recent work on uncertainty quantification for wearable "
                "biomedical signals, including a 2025 systematic comparison of "
                "eight UQ techniques for photoplethysmography that explicitly "
                "evaluates conformal calibration on clinically relevant "
                "prediction tasks. Direct methodological overlap with TinyDCS.",
        },
        {
            "n": 2,
            "name": "Esther Rodriguez-Villegas, PhD, FREng",
            "affil": "Professor of Low Power Electronics, Department of "
                     "Electrical and Electronic Engineering, Imperial College "
                     "London, UK",
            "email": "e.rodriguez@imperial.ac.uk",
            "orcid": "(see Imperial profile)",
            "rationale":
                "Leads research on TinyML and embedded ML for healthcare "
                "wearables; her IEEE Access review surveys exactly the "
                "MCU-class deployment regime that TinyDCS targets "
                "(sub-100 KB ONNX, microsecond-class latency).",
        },
        {
            "n": 3,
            "name": "Francesco Conti, PhD",
            "affil": "Associate Professor, Department of Electrical, "
                     "Electronic, and Information Engineering ‘Guglielmo "
                     "Marconi’ (DEI), University of Bologna, Italy",
            "email": "f.conti@unibo.it",
            "orcid": "(see DEI / Google Scholar)",
            "rationale":
                "TinyML hardware and software for ultra-low-power "
                "systems-on-chip (PULP Platform), with recent biomedical "
                "applications. His 2026 BioTrain paper covers the same "
                "sub-MB sub-50 mW deployment regime as TinyDCS.",
        },
        {
            "n": 4,
            "name": "Allan Peter Engsig-Karup, PhD",
            "affil": "Associate Professor in Scientific Computing, Department "
                     "of Applied Mathematics and Computer Science (DTU "
                     "Compute), Technical University of Denmark, Kongens "
                     "Lyngby, Denmark",
            "email": "apek@dtu.dk",
            "orcid": "0000-0001-8626-1575",
            "rationale":
                "Recent work applying conformal prediction with Venn-ABERS "
                "calibration to clinical bacterial-infection-focus prediction "
                "(gradient-boosting risk classifiers + conformal risk "
                "control) — direct methodological parallel to TinyDCS.",
        },
        {
            "n": 5,
            "name": "Nils Strodthoff, PhD",
            "affil": "Professor of eHealth (Interpretable and Explainable "
                     "Learning Algorithms), Department of Health Services "
                     "Research, Carl von Ossietzky Universität Oldenburg, "
                     "Germany",
            "email": "nils.strodthoff@uni-oldenburg.de",
            "orcid": "(see UOL profile)",
            "rationale":
                "Leads the AI4Health group at U Oldenburg; first author of "
                "Deep Learning for ECG Analysis: Benchmarks and Insights from "
                "PTB-XL (IEEE JBHI 2021). Independent voice on whether "
                "TinyDCS’s UQ design choices are state-of-the-art.",
        },
    ]

    for r in reviewers:
        _add_para(doc, f"Reviewer {r['n']} — {r['name']}",
                  bold=True, space_after=2)
        for label, val in (("Affiliation", r["affil"]),
                           ("Email", r["email"]),
                           ("ORCID", r["orcid"]),
                           ("Rationale", r["rationale"])):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{label}: ")
            run.bold = True
            _add_inline(p, val)
        doc.add_paragraph()

    out = OUT / "Suggested Reviewers.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ---------------------------------------------------------------------------
# Stage figures
# ---------------------------------------------------------------------------

def stage_figures() -> None:
    """Copy figure PDFs into the submission folder with portal-friendly names."""
    fig_map = {
        "fig1_reliability_diagram.pdf": "Figure 1.pdf",
        "fig2_per_band_coverage.pdf": "Figure 2.pdf",
        "fig3_size_vs_accuracy.pdf": "Figure 3.pdf",
        "fig4_personalization_info_gain.pdf": "Figure 4.pdf",
        "fig5_architecture.pdf": "Figure 5.pdf",
    }
    for src_name, dst_name in fig_map.items():
        src = FIG_SRC / src_name
        dst = OUT / dst_name
        shutil.copyfile(src, dst)
        print(f"Copied {src.name} -> {dst.name}")


# ---------------------------------------------------------------------------
# Build everything
# ---------------------------------------------------------------------------

def main() -> None:
    build_manuscript()
    build_cover_letter()
    build_highlights()
    build_coi()
    build_human_animal()
    build_credit()
    build_reviewers()
    stage_figures()
    print("\nAll deliverables in:", OUT)


if __name__ == "__main__":
    main()
